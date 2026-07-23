from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt


ROOT = Path(__file__).resolve().parents[2]
PREVIEW = ROOT / "tmp" / "pdfs" / "depliant-preview-1.png"
OUT = ROOT / "output" / "pdf" / "depliant-fermeture-poche-falaise-chambois-pages.docx"


def main():
    if not PREVIEW.exists():
        raise FileNotFoundError(f"Missing rendered page: {PREVIEW}")

    OUT.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(0)
    section.bottom_margin = Mm(0)
    section.left_margin = Mm(0)
    section.right_margin = Mm(0)
    section.header_distance = Mm(0)
    section.footer_distance = Mm(0)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(1)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run()
    run.add_picture(str(PREVIEW), width=Mm(210))

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
